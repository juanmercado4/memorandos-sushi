import os
import json
import re
import uuid
from datetime import date
from io import BytesIO

from flask import Flask, render_template, request, jsonify, send_file, session
import anthropic

# Server-side store to avoid Flask cookie 4 KB limit
_doc_store: dict = {}

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "memorandos-sushi-break-2024")

EMPRESA = "SUSHI BREAK SAS"
NIT = "901359641-1"
REPRESENTANTE = "Juan Manuel Mercado"
CARGO_REP = "Representante Legal"

SEDES = ["NOVENA", "NORTE", "OESTE", "V LILI", "JAMUNDI", "PALMIRA", "CP", "MARBELLA"]

TIPOS_DOCUMENTO = {
    "llamado":     "MEMORANDO DE LLAMADO DE ATENCIÓN",
    "descargos":   "CITACIÓN A DILIGENCIA DE DESCARGOS",
    "sancion":     "MEMORANDO DE SANCIÓN DISCIPLINARIA",
    "terminacion": "CARTA DE TERMINACIÓN DE CONTRATO",
}

SYSTEM_PROMPT = """Eres un experto en derecho laboral colombiano al servicio de SUSHI BREAK SAS (NIT 901359641-1), empresa del sector gastronómico con sedes en Cali y área metropolitana. Tu tarea es redactar documentos disciplinarios laborales formales, citando el Reglamento Interno de Trabajo (RIT) de Sushi Break SAS vigente desde el 01-06-2026 y el Código Sustantivo del Trabajo (CST).

=== EMPRESA ===
Razón social: SUSHI BREAK SAS
NIT: 901359641-1
Representante Legal: Juan Manuel Mercado
Sedes: NOVENA, NORTE, OESTE, V LILI, JAMUNDI, PALMIRA, CP, MARBELLA

=== CLASIFICACIÓN DE FALTAS (RIT Sushi Break - Art. 68) ===

FALTAS LEVES (Art. 68.1 RIT) — Solo requieren llamado de atención preventivo:
1. Retardo ocasional en la hora de entrada que no supere 15 minutos
2. Incumplimiento leve de obligaciones contractuales o reglamentarias sin afectación relevante
3. Incumplimiento ocasional de normas internas de comportamiento
4. Incumplimiento de normas de presentación personal o uso del uniforme
5. No mantener el orden y aseo en el puesto de trabajo

FALTAS MEDIAS (Art. 68.2 RIT) — Requieren procedimiento disciplinario, pueden generar suspensión hasta 8 días:
1. Retardos reiterados (más de 2 veces en el mismo mes o mayores a 15 min sin justificación)
2. Cambio de turnos sin autorización del jefe inmediato
3. Incumplimiento reiterado de obligaciones contractuales
4. Uso de celular o dispositivos personales durante la jornada afectando atención a clientes
5. Ausentarse del puesto de trabajo sin previa autorización
6. Escuchar música durante el turno afectando comunicación o servicio
7. Presentación incorrecta de platos de forma reiterada
8. Actitudes negligentes frente a clientes
9. Errores en recetas por descuido o incumplimiento de instrucciones
10. Incumplir protocolos de apertura de cocina (inventarios, chequeos)
11. Errores en caja por descuido y/o descuadres
12. No registrar correctamente ventas o movimientos
13. No seguir protocolos de cierre o apertura de caja
14. No reportar novedades operativas importantes

FALTAS GRAVES (Art. 68.3 RIT) — Pueden dar lugar a terminación con justa causa:
CALIDAD Y COCINA: Alterar calidad del producto intencionalmente, incumplir recetas afectando la esencia del producto, preparar alimentos en condiciones de riesgo para la salud, contaminar alimentos por negligencia grave, introducir o consumir alimentos de origen animal en zonas veganas.
MANEJO DE DINERO: Apropiarse de dinero de caja, realizar "jineteo" de dinero, prestar dinero de caja sin autorización, alterar registros de ventas o facturación, omitir ingresos o manipular información contable, descuadres reiterativos de caja, recibir pagos fuera de canales autorizados.
CONDUCTA Y DISCIPLINA: Abandonar el puesto de trabajo, negarse a cumplir instrucciones, reincidir en faltas medias, uso del celular con abandono del puesto.
ÉTICA Y HONESTIDAD: Fraude, deshonestidad o engaño, falsificar documentos, presentar información falsa, revelar información confidencial o recetas, realizar actos de competencia desleal.
RESPETO Y CONVIVENCIA: Tratar irrespetuosamente a clientes o compañeros, agresiones físicas o verbales, acoso laboral o sexual (Ley 1010/2006 y Ley 2365/2024), conductas discriminatorias.
SEGURIDAD: Presentarse bajo efectos de alcohol o sustancias psicoactivas, consumir o introducir sustancias prohibidas, incumplir normas SG-SST poniendo en riesgo la seguridad.

=== CIRCUNSTANCIAS AGRAVANTES (Art. 70 RIT) ===
- Reincidencia en la comisión de faltas laborales (Art. 72 RIT: dentro de 12 meses)
- Intención de causar perjuicio a la Empresa
- Abuso de la confianza depositada
- Afectación grave a clientes o terceros
- Comisión de la falta utilizando engaño o fraude

=== CIRCUNSTANCIAS ATENUANTES (Art. 71 RIT) ===
- No tener antecedentes disciplinarios
- Reconocer voluntariamente la falta
- Haber procurado evitar o disminuir el daño
- Antigüedad y buen desempeño previo

=== MATRIZ DISCIPLINARIA (Art. 73 RIT) ===
Falta leve → Llamado de atención preventivo (no requiere procedimiento disciplinario)
Falta media → Llamado de atención fuerte / Suspensión hasta 8 días (requiere procedimiento)
Falta grave → Suspensión hasta 2 meses / Terminación con justa causa (requiere procedimiento)

=== PROCEDIMIENTO DISCIPLINARIO (Arts. 75-84 RIT) ===
- Art. 75: Garantías del debido proceso, derecho de defensa y presunción de inocencia
- Art. 77: Apertura del proceso con comunicación escrita al trabajador
- Art. 78: El trabajador tiene mínimo 5 días hábiles para presentar descargos
- Art. 79: Empresa evalúa pruebas y adopta decisión motivada por escrito
- Art. 80: El trabajador puede solicitar revisión dentro de 3 días hábiles
- Art. 81: Terminación de contrato por justa causa garantiza debido proceso previo
- Art. 82: Llamados de atención preventivos NO son sanción disciplinaria
- Art. 83: Sanciones disciplinarias: (1) Llamado de atención fuerte, (2) Suspensión del contrato
- Art. 84: Sanción impuesta sin procedimiento no produce efecto alguno

=== OBLIGACIONES DE LOS TRABAJADORES (Art. 64 RIT) ===
- Cumplir políticas, manuales, reglamentos y directrices de la Empresa
- Cumplir código de vestimenta y normas de presentación personal e imagen corporativa
- Guardar confidencialidad de recetas, bases de datos de clientes y estrategias comerciales
- Usar y administrar recursos y bienes de la Empresa de manera honesta y eficiente
- Informar inmediatamente cualquier situación anómala con bienes de la Empresa

=== PROHIBICIONES A LOS TRABAJADORES (Art. 66 RIT — selección relevante) ===
- Sustraer, apropiarse o retener bienes de la Empresa sin autorización (numeral 7)
- Presentarse al trabajo en estado de embriaguez o bajo efectos de sustancias (numeral 31)
- Ingerir o consumir alcohol o sustancias psicoactivas durante la jornada (numeral 32)
- Faltar al trabajo sin justa causa o sin permiso (numeral 36)
- Registrar marcación de control de horario de otro trabajador (numeral 40)
- Dormir durante la jornada laboral (numeral 44)
- Revelar información que afecte los intereses de la Empresa (numeral 21)
- Realizar cobros no autorizados a clientes (numeral relacionado)
- Conductas de acoso laboral o sexual (numeral 68)

=== INSTRUCCIONES DE REDACCIÓN ===
1. El usuario ya eligió el tipo de documento — redáctalo exactamente como lo solicitó
2. Cita SIEMPRE artículos del RIT de Sushi Break SAS Y del CST
3. Usa tono formal, legal y profesional
4. El documento debe mencionar: empresa, NIT, nombre del trabajador, cargo, sede, fecha, hechos, normas infringidas y decisión/citación
5. Para LLAMADO DE ATENCIÓN: tono correctivo pero pedagógico, cita Art. 82 y 68.1 RIT
6. Para CITACIÓN A DESCARGOS: informar hechos investigados, citar Art. 77-78 RIT, indicar plazo de 5 días hábiles para descargos
7. Para SANCIÓN DISCIPLINARIA: especificar días de suspensión, citar Art. 83 RIT, Art. 112-113 CST
8. Para TERMINACIÓN DE CONTRATO: citar causal específica del Art. 62 CST y Art. 81 RIT, indicar que se garantizó el debido proceso
9. Si hay reincidencia, citar expresamente el Art. 70 y 72 RIT como circunstancia agravante

RESPONDE ÚNICAMENTE con un JSON válido con esta estructura exacta:
{
  "tipo_falta": "LEVE|MEDIA|GRAVE",
  "clasificacion": "descripción breve de la situación y norma infringida",
  "articulos": ["Art. XX RIT Sushi Break - descripción", "Art. XX CST - descripción", ...],
  "recomendacion": "observación legal relevante para el empleador",
  "tipo_documento": "nombre del documento tal como lo indicó el usuario",
  "documento": {
    "asunto": "texto del asunto",
    "cuerpo": "texto completo del cuerpo del documento con saltos de línea \\n"
  }
}"""


@app.route("/")
def index():
    today = date.today().strftime("%Y-%m-%d")
    return render_template("index.html", sedes=SEDES, today=today)


@app.route("/analizar", methods=["POST"])
def analizar():
    data = request.get_json()
    nombre = data.get("nombre", "").strip()
    cargo = data.get("cargo", "").strip()
    sede = data.get("sede", "").strip()
    fecha = data.get("fecha", "").strip()
    reincidente = data.get("reincidente", "no")
    descripcion = data.get("descripcion", "").strip()
    tipo_key = data.get("tipo_documento", "llamado")

    if not all([nombre, cargo, sede, fecha, descripcion]):
        return jsonify({"error": "Todos los campos son obligatorios"}), 400

    tipo_doc_label = TIPOS_DOCUMENTO.get(tipo_key, "MEMORANDO DE LLAMADO DE ATENCIÓN")

    user_message = f"""Genera el siguiente documento disciplinario:

TIPO DE DOCUMENTO SOLICITADO: {tipo_doc_label}

DATOS DEL EMPLEADO:
- Nombre: {nombre}
- Cargo: {cargo}
- Sede: {sede}
- Fecha del documento: {fecha}
- ¿Es reincidente?: {reincidente.upper()}

DESCRIPCIÓN DE LA SITUACIÓN:
{descripcion}

Redacta el documento completo del tipo indicado ({tipo_doc_label}). Usa los artículos del CST correspondientes a ese tipo de documento."""

    try:
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
        message = client.messages.create(
            model="claude-haiku-4-5",
            max_tokens=2048,
            system=SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_message}]
        )
        raw = message.content[0].text.strip()

        # Extract JSON even if wrapped in markdown code block
        json_match = re.search(r'\{[\s\S]*\}', raw)
        if not json_match:
            return jsonify({"error": "Respuesta inválida del modelo"}), 500
        resultado = json.loads(json_match.group())

        # Store in server-side dict (avoids Flask cookie 4 KB limit)
        doc_id = str(uuid.uuid4())
        _doc_store[doc_id] = {
            "resultado": resultado,
            "empleado": {
                "nombre": nombre,
                "cargo": cargo,
                "sede": sede,
                "fecha": fecha,
                "reincidente": reincidente,
            },
        }
        # Keep only the last 50 docs to avoid unbounded memory growth
        if len(_doc_store) > 50:
            oldest = next(iter(_doc_store))
            del _doc_store[oldest]

        session["doc_id"] = doc_id
        return jsonify(resultado)

    except json.JSONDecodeError:
        return jsonify({"error": "No se pudo parsear la respuesta del modelo"}), 500
    except anthropic.APIError as e:
        return jsonify({"error": f"Error de API: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Error inesperado: {str(e)}"}), 500


def _format_date_spanish(fecha_str):
    months = {
        "01": "enero", "02": "febrero", "03": "marzo", "04": "abril",
        "05": "mayo", "06": "junio", "07": "julio", "08": "agosto",
        "09": "septiembre", "10": "octubre", "11": "noviembre", "12": "diciembre"
    }
    try:
        parts = fecha_str.split("-")
        return f"{int(parts[2])} de {months[parts[1]]} de {parts[0]}"
    except Exception:
        return fecha_str


@app.route("/generar_pdf")
def generar_pdf():
    doc_id = session.get("doc_id")
    entry = _doc_store.get(doc_id) if doc_id else None
    if not entry:
        return "No hay documento generado. Vuelve a analizar la situación.", 400
    resultado = entry["resultado"]
    empleado = entry["empleado"]

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2.5 * cm
    )

    styles = getSampleStyleSheet()
    story = []

    # Header style
    header_style = ParagraphStyle(
        "Header",
        parent=styles["Normal"],
        fontSize=14,
        fontName="Helvetica-Bold",
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    sub_header_style = ParagraphStyle(
        "SubHeader",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        textColor=colors.white,
        alignment=TA_CENTER,
        spaceAfter=2,
    )
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Normal"],
        fontSize=12,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
        spaceBefore=14,
        spaceAfter=14,
        textColor=colors.HexColor("#1a1a2e"),
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontSize=9,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#444444"),
        spaceAfter=2,
    )
    value_style = ParagraphStyle(
        "Value",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        textColor=colors.black,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        alignment=TA_JUSTIFY,
        spaceAfter=10,
        leading=16,
    )
    sign_style = ParagraphStyle(
        "Sign",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
    )

    # Dark header block
    header_data = [
        [Paragraph(EMPRESA, header_style)],
        [Paragraph(f"NIT: {NIT}", sub_header_style)],
        [Paragraph(resultado["tipo_documento"], sub_header_style)],
    ]
    header_table = Table(header_data, colWidths=[16 * cm])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#1a1a2e")),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 0.4 * cm))

    # Employee info table
    fecha_es = _format_date_spanish(empleado["fecha"])
    info_data = [
        [Paragraph("PARA:", label_style), Paragraph(f"{empleado['nombre']} — {empleado['cargo']}", value_style)],
        [Paragraph("SEDE:", label_style), Paragraph(empleado["sede"], value_style)],
        [Paragraph("DE:", label_style), Paragraph(f"{REPRESENTANTE} — {CARGO_REP}", value_style)],
        [Paragraph("ASUNTO:", label_style), Paragraph(resultado["documento"]["asunto"], value_style)],
        [Paragraph("FECHA:", label_style), Paragraph(fecha_es, value_style)],
    ]
    info_table = Table(info_data, colWidths=[2.5 * cm, 13.5 * cm])
    info_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LINEBELOW", (0, -1), (-1, -1), 0.5, colors.HexColor("#cccccc")),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 0.4 * cm))

    # Body paragraphs
    for line in resultado["documento"]["cuerpo"].split("\n"):
        line = line.strip()
        if line:
            story.append(Paragraph(line, body_style))

    story.append(Spacer(1, 1.5 * cm))
    story.append(HRFlowable(width="40%", thickness=1, color=colors.HexColor("#1a1a2e"), hAlign="CENTER"))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(REPRESENTANTE, sign_style))
    story.append(Paragraph(CARGO_REP, ParagraphStyle("sc", parent=sign_style, fontName="Helvetica", fontSize=9)))
    story.append(Paragraph(EMPRESA, ParagraphStyle("sc2", parent=sign_style, fontName="Helvetica", fontSize=9)))

    doc.build(story)
    buffer.seek(0)

    filename = f"memorando_{empleado['nombre'].replace(' ', '_')}.pdf"
    return send_file(buffer, mimetype="application/pdf",
                     as_attachment=True, download_name=filename)


@app.route("/generar_word")
def generar_word():
    doc_id = session.get("doc_id")
    entry = _doc_store.get(doc_id) if doc_id else None
    if not entry:
        return "No hay documento generado. Vuelve a analizar la situación.", 400
    resultado = entry["resultado"]
    empleado = entry["empleado"]

    document = Document()

    # Page margins
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # Header table
    header_table = document.add_table(rows=3, cols=1)
    header_table.style = "Table Grid"
    rows_data = [EMPRESA, f"NIT: {NIT}", resultado["tipo_documento"]]
    font_sizes = [14, 10, 10]
    for i, (text, fsize) in enumerate(zip(rows_data, font_sizes)):
        cell = header_table.rows[i].cells[0]
        set_cell_bg(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = (i == 0)
        run.font.size = Pt(fsize)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    document.add_paragraph()

    # Employee info
    fecha_es = _format_date_spanish(empleado["fecha"])
    fields = [
        ("PARA:", f"{empleado['nombre']} — {empleado['cargo']}"),
        ("SEDE:", empleado["sede"]),
        ("DE:", f"{REPRESENTANTE} — {CARGO_REP}"),
        ("ASUNTO:", resultado["documento"]["asunto"]),
        ("FECHA:", fecha_es),
    ]
    info_table = document.add_table(rows=len(fields), cols=2)
    info_table.style = "Table Grid"
    col_widths = [Cm(2.8), Cm(13.2)]
    for row_idx, (label, value) in enumerate(fields):
        row = info_table.rows[row_idx]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        lp = row.cells[0].paragraphs[0]
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)

    document.add_paragraph()

    # Body
    for line in resultado["documento"]["cuerpo"].split("\n"):
        line = line.strip()
        if line:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line)
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(8)

    # Signature
    document.add_paragraph()
    document.add_paragraph()
    sig_p = document.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = sig_p.add_run("_" * 35)
    sig_run.font.size = Pt(10)

    for line in [REPRESENTANTE, CARGO_REP, EMPRESA]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = (line == REPRESENTANTE)
        run.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"memorando_{empleado['nombre'].replace(' ', '_')}.docx"
    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


if __name__ == "__main__":
    app.run(debug=True)
